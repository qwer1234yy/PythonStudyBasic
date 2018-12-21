import SMART.Smart_urls as base_url
import SMART.Smart_com_acts as Act
from docx import Document
from docx.shared import Inches
import os, time
from selenium.webdriver.common.by import By
import SMART.Smart_common_settings as settings


def login_(driver):

    driver.get(base_url.base_url)

    login_ele = driver.find_element_by_id('txtLoginid')
    password_ele = driver.find_element_by_id('txtPassword')
    login_ele_btn = driver.find_element_by_id('btnloginText')

    login_ele.clear()
    login_ele.send_keys('admin')
    password_ele.clear()
    password_ele.send_keys('Admin')
    login_ele_btn.click()
    # wait for login completed
    Act.wait_invisibility_of_element_located(driver)
    Act.wait_document_completed(driver)
    Act.wait_presence_element(driver, 'aModuleSIP101')
    Act.wait_element_clickable(driver, By.ID, 'aModuleSIP101')
    Act.wait_element_clickable(driver, By.ID, 'aSlideMenuSelModuleOP101')
    time.sleep(1)


def read_custom_fields():
    fields = open('custom_search_fields.txt', 'r')
    lines = fields.readlines()
    print(lines.__len__())
    for line in lines:
            print('----------' + line.strip() + '-------')
    fields.close()

def read_file_as_list(file_name):
    print('read_as_list' + file_name)
    fields_f = open(file_name, 'r')
    lines = fields_f.readlines()
    fields = []

    print(lines.__len__())
    for line in lines:
        print(line.strip())
        if not (line.strip().startswith('-')):
            fields.append(line)
    fields_f.close()
    return fields
def go_to_MS_report(driver):
    print(go_to_MS_report)
    MS_access = driver.find_element_by_id('aSlideMenuSelModuleIP101')
    MS_access.click()

    MS_report_glossary_access = driver.find_element_by_id('aModuleReportGlossary')
    MS_report_glossary_access.click()


def read_file_report_names_as_list(enterprise_or_standard):
    file_name = '../Smart_resources/op_report_names_us300299.txt'
    print(os.getcwd())
    print('read_as_list---' + file_name)
    report_name_f = open(file_name, 'r')
    lines = report_name_f.readlines()
    report_names = []
    report_name_enterprise = []
    enterprise_case_num = []
    report_name_standard = []
    standard_case_num = []

    report_name_enterprise_dic = {}
    report_name_standard_dic = {}

    print(lines.__len__())
    for line in lines:
        print(line.strip())
        if not (line.strip().startswith('-')):
            if '-standard' in line:
                print('-----------IP standard report-------------')
                enterprise_case_num.append(line.strip()[-6:])
                report_name_standard.append(line.strip()[0:-16])
                report_name_standard_dic[line.strip()[-6:]] = line.strip()[0:-16]

            elif '-enterprise' in line:
                print('-----------IP enterprise report-------------')
                standard_case_num.append(line.strip()[-6:])
                report_name_enterprise.append(line.strip()[0:-18])
                report_name_enterprise_dic[line.strip()[-6:]] = line.strip()[0:-18]
            else:
                print('-----------MS report-------------')
    report_name_f.close()
    if enterprise_or_standard == 'standard':
        return report_name_standard_dic
    elif enterprise_or_standard == 'enterprise':
        return report_name_enterprise_dic


def read_file_as_list_report_US299725(enterprise_or_standard):
    file_name = '../reports_names.txt'
    print('read_as_list' + file_name)
    fields_f = open(file_name, 'r')
    lines = fields_f.readlines()
    fields = []
    fields_enterprise = []
    fields_enterprise_num = []
    fields_standard = []
    fields_standard_num = []

    fields_enterprise_dic = {}
    fields_standard_dic = {}

    print(lines.__len__())
    for line in lines:
        print(line.strip())
        if not(line.strip().startswith('-')):
            if '-IPST' in line:
                print('-----------IP standard report-------------')
                fields_enterprise_num.append(line.strip()[-6:])
                fields_standard.append(line.strip()[0:-12])
                fields_standard_dic[line.strip()[-6:]] = line.strip()[0:-12]

            elif '-IPEN' in line:
                print('-----------IP enterprise report-------------')
                fields_standard_num.append(line.strip()[-6:])
                fields_enterprise.append(line.strip()[0:-12])
                fields_enterprise_dic[line.strip()[-6:]] = line.strip()[0:-12]
            else:
                print('-----------MS report-------------')
    fields_f.close()
    if enterprise_or_standard == 'standard':
        return fields_standard_dic
    elif enterprise_or_standard == 'enterprise':
        return fields_enterprise_dic


def write_test_result_as_docx(report_name, test_case):

    if os.path.exists('../results/'+settings.test_result_file_name):
        document = Document('../results/'+settings.test_result_file_name)
    else:
        document = Document()
    document.add_heading(test_case, level=1)
    document.add_picture(report_name, width=Inches(6.43))
    document.save('../results/'+settings.test_result_file_name)


def write_test_result_as_docx_com(screenshot_path, heading_value):

    if os.path.exists('../results/'+settings.test_result_file_name):
        document = Document('../results/'+settings.test_result_file_name)
    else:
        document = Document()
    document.add_heading(heading_value, level=1)
    document.add_picture(screenshot_path, width=Inches(6.43))
    document.save('../results/'+settings.test_result_file_name)

def file_exists_delete(file_path):
    if os.path.exists(file_path):
        print('------------file_exists_delete-----------')
        os.remove(file_path)


def screen_shot_report_us_test_result(driver, report_name, report_test_case_number):

    test_case_number = 'TC ' + str(report_test_case_number) + '-' + report_name
    if report_name == 'Top 50 CC/MCC Diagnoses':
        picture_path = 'pictures/Top 50 CC_MCC Diagnoses.png'
    else:
        picture_path = '../pictures/' + report_name + '.png'
    file_exists_delete(picture_path)
    driver.save_screenshot(picture_path)
    write_test_result_as_docx(picture_path, test_case_number)