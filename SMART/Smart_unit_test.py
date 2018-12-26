import unittest,time,os
import SMART.Smart_commons as SC
from selenium import webdriver
import SMART.Smart_com_acts as ACT
from docx import Document
from docx.shared import Inches
import SMART.KeyboardKeys as KEY
import win32com, win32api
from selenium.webdriver.common.desired_capabilities import DesiredCapabilities
from winreg import *
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.action_chains import ActionChains
import SMART.Smart_tools.tools as tools
import SMART.Smart_tools.report_tools as report_tools


class MyTestCase(unittest.TestCase):

    def setUp(self):
        print('setup')
    def tearDown(self):
        print('tearDown')

    def test_report_tool(self):
        names = report_tools.read_report_names('standard')
        print(names)


    def test_tools(self):
        # tools.spide_write_to_txt('sssss','Smart_resources/cs_fields.txt')

        field1 = {'id': '12213', 'operators': 'ewe', 'type': 'wwe', 'default': 'wewewe'}
        field2 = {'id': '12213', 'operators': 'ewe', 'type': 'wwe', 'default': 'wewewe'}
        field3 = {'id': '12213', 'operators': 'ewe', 'type': 'wwe', 'default': 'wewewe'}
        field4 = {'id': '12213', 'operators': 'ewe', 'type': 'wwe', 'default': 'wewewe'}
        field5 = {'id': '12213', 'operators': 'ewe', 'type': 'wwe', 'default': 'wewewe'}

        dics_list = [field1, field2, field3, field4, field5]
        file_patt = 'Smart_resources/cs_fields.xml'
        tools.write_dics_list_to_xml('Fields', 'Field', dics_list, file_patt)

        person1 = {'name': 'Jon', 'gender': 'man', 'id':'JOn'}
        person2 = {'name': 'Jon', 'gender': 'man', 'id':'JOn'}
        person3 = {'name': 'Jon', 'gender': 'man', 'id':'JOn'}
        person4 = {'name': 'Jon', 'gender': 'man', 'id':'JOn'}

        person = [person1, person2, person3, person4]
        person_file_patt = 'Smart_resources/person.xml'
        tools.write_dics_list_to_xml('people','person',person, person_file_patt)

    def test_os(self):
        print('----------'+os.getcwd())
        exist = os._exists(os.getcwd() + '/Smart_resources/op_report_names_us300299.txt')
        print(exist)



    def test_test(self):
        file_name = '../SMART/Smart_resources/op_report_names_us300299.txt'
        print('read_as_list---' + file_name)
        report_name_f = open(file_name, 'r')
        lines = report_name_f.readlines()
        for l in lines:
            print(l)


    def test_write_result(self):
        # Flag Comparison.png
        # 'Medicare Medical Necessity Flags.png'
        SC.write_test_result_as_docx('pictures/1.PNG', 'TC 123232 -xxxxx')

    def test_wait(self):
        driver = webdriver.Firefox()
        SC.login_(driver)

        # Click "inpatient"
        inpatient = driver.find_element_by_id('aModuleSIP101')
        inpatient.click()

        # Click "report"
        report = driver.find_element_by_id('aIPModuleWorkplans')
        report.click()

    def test_txt(self):
        fields_f = open('custom_search_fields.txt','r')
        lines = fields_f.readlines()
        fields = []
        print(lines.__len__())
        for line in lines:
             print(line.strip())
             fields.append(line.strip())
        fields_f.close()

    def test_get_default_fields(self):
        print('test_get_default_fields')

    def test_read_as_list(self):
        # custom_search_fields.txt
        items = SC.read_file_as_list('reports_names.txt')

        for item in items:
            print(item)
    def test_read_as_list_reports(self):
        # standard enterprise
        reports_dic_enterprise = SC.read_file_as_list_report_US299725('enterprise')
        reports_dic_standard = SC.read_file_as_list_report_US299725('standard')

        enterprise = list(reports_dic_enterprise.values())

        for i in enterprise:
            print(i)


    def test_python_docx(self):
        document = Document()

        document.add_heading('Document Title', 0)

        p = document.add_paragraph('A plain paragraph having some ')
        p.add_run('bold').bold = True
        p.add_run(' and some ')
        p.add_run('italic.').italic = True

        document.add_heading('Heading, level 1', level=1)
        document.add_paragraph('Intense quote', style='Intense Quote')

        document.add_paragraph(
            'first item in unordered list', style='List Bullet'
        )
        document.add_paragraph(
            'first item in ordered list', style='List Number'
        )

        document.add_picture('pictures/DRG Clinical Profile.png', width=Inches(1.25))

        records = (
            (3, '101', 'Spam'),
            (7, '422', 'Eggs'),
            (4, '631', 'Spam, spam, eggs, and spam')
        )

        table = document.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Qty'
        hdr_cells[1].text = 'Id'
        hdr_cells[2].text = 'Desc'
        for qty, id, desc in records:
            row_cells = table.add_row().cells
            row_cells[0].text = str(qty)
            row_cells[1].text = id
            row_cells[2].text = desc

        document.add_page_break()

        document.save('demo.docx')
    def test_docx_for_299725(self):

        pictures = 'pictures/DRG Clinical Profile.png'
        headings = 'Test Case 123456'
        SC.write_test_result_as_docx(pictures,headings)

    def test_kayBoeard(self):
        print('test')

    def test_compatibility(self):

        reports_dic_standard_all = SC.read_file_as_list_report_US299725('standard')
        reports_dic_enterprise_all = SC.read_file_as_list_report_US299725('enterprise')
        print(reports_dic_standard_all.__len__())
        print(reports_dic_enterprise_all.__len__())

        reports_dic_standard_keys = list(reports_dic_standard_all.keys())
        reports_dic_standard_values = list(reports_dic_standard_all.values())

        reports_dic_standard = {}
        for dic in range(reports_dic_standard_all.__len__()):

            reports_dic_standard[reports_dic_standard_keys[dic]] = reports_dic_standard_values[dic]

        for i in reports_dic_standard.items():
            print(i)
    def test_selenium_ie(self):

        # try:
        #     keyVal = r'Software\Microsoft\Windows\CurrentVersion\Internet Settings\Zones\1'
        #     key = OpenKey(HKEY_CURRENT_USER, keyVal, 0, KEY_ALL_ACCESS)
        #     SetValueEx(key, "2500", 0, REG_DWORD, 0)
        #     print("enabled protected mode")
        # except Exception:
        #     print("failed to enable protected mode")

        caps = DesiredCapabilities.INTERNETEXPLORER
        caps["platform"] = "WINDOWS"
        caps["browserName"] = "internet explorer"
        caps["requireWindowFocus"] = True
        # IE_path = 'C:\Program Files\Internet Explorer\iexplore.exe'
        # os.environ['webdriver.ie.driver'] = IE_path
        # driver = webdriver.Ie(capabilities=caps,executable_path=IE_path)
        driver = webdriver.Ie(capabilities=caps)
        driver.get('https://www.baidu.com/')
        ACT.wait_document_completed(driver)
        time.sleep(10)
        search_box = driver.find_element_by_id('kw')
        search_box.send_keys('python')
    def test_Chrome_options(self):
        print('test_Chrome')
        caps = DesiredCapabilities.CHROME

        # chrome_options = Options()
        # chrome_options.add_argument("--disable-extensions")
        # driver = webdriver.Chrome(chrome_options=chrome_options)

        # test-type no-sandbox disable-extensions start-maximized --js-flags=--expose-gc disable-plugins
        # --enable-precise-memory-info --disable-popup-blocking --disable-default-apps test-type=browser disable-infobars
        # chrome_options = webdriver.ChromeOptions()
        # chrome_options.add_argument('--disable-extensions')
        # chrome_options.add_argument('--no-sandbox')
        # driver = webdriver.Chrome(chrome_options=chrome_options)

        capabilities = {
            'browserName': 'chrome',
            'chromeOptions': {
                'useAutomationExtension': False,
                'forceDevToolsScreenshot': True,
                'args': ['--start-maximized', '--disable-infobars']
            }
        }
        driver = webdriver.Chrome(desired_capabilities=capabilities)

        driver.get('http://strzw058051/SMARTSolutions/')
    def test_chrome_taps(self):
        print('test_chrome_taps')
        driver = webdriver.Firefox()
        driver.maximize_window()
        driver.get("http://www.baidu.com/")

        ActionChains(driver).key_down(Keys.CONTROL).send_keys('t').perform()

        # open tab
        # driver.actions().keyDown(Keys.CONTROL).sendKeys('t').perform()
        driver.find_element_by_tag_name('body').send_keys(Keys.CONTROL + 't')
        # You can use (Keys.CONTROL + 't') on other OSs

        # Load a page
        driver.get('http://stackoverflow.com/')
        # Make the tests...

        # close the tab
        # (Keys.CONTROL + 'w') on other OSs.
        # driver.actions().keyDown(Keys.CONTROL).sendKeys('w').perform()
        driver.find_element_by_tag_name('body').send_keys(Keys.CONTROL + 'w')


    def test_read_file_report_names_as_list(self):
        reports_name_cases = SC.read_file_report_names_as_list('enterprise')

        report_names = []
        report_case_num = []

        for i in reports_name_cases:
            report_names.append(reports_name_cases[i])
            report_case_num.append(i)

        print(report_names)
        print(report_case_num)
        #
        # items = SC.read_file_report_names_as_list('standard')
        # for i in items:
        #     print(i)
        #     print(items[i])


if __name__ == '__main__':
    unittest.main()
