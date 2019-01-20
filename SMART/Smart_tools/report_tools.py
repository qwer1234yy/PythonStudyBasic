from selenium.webdriver.common.by import By
import SMART.Smart_common.Smart_actions as Smart_actions
import SMART.Smart_tools.report_tools as report_tools
import SMART.Smart_pages.Smart_reports as reports
import SMART.Smart_common_settings as settings
import os
from SMART.Smart_tools import tools
import SMART.Smart_com_acts as ACT
from docx import Document
from docx.shared import Inches


def report_continue(driver, report_name):
    pass


def timeout_handle(driver, report_name, timeout):
    report_tools.take_screenshot(driver, report_name + 'out of memory')
    report_load_timeout = 'report_load_timeout' + timeout.__str__()
    print(report_load_timeout)
    reports.reports.close_report_view_window(driver)
    # 加一个filter
    # set up filters
    # click "+" icon
    Smart_actions.wait_element_clickable(driver, By.ID, 'btnInsertClause')
    add_icon = driver.find_element_by_id('btnInsertClause')
    add_icon.click()

    # add filters besides the default ones
    driver.find_element_by_id('customsearch-grid-div')
    fields = driver.find_elements_by_xpath(
        './/div[@id="customsearch-grid-div"]/div/div[@class="k-grid-content"]/table/tbody/tr')

    # click the last row of the CS
    field_last = driver.find_element_by_xpath(
        './/div[@id="customsearch-grid-div"]/div/div[@class="k-grid-content"]/table/tbody/tr[last()]/td[3]/span')
    field_last.click()

    fields_lis = driver.find_elements_by_xpath('//div[@class="k-animation-container"][last()]/div/ul/li')
    print(fields_lis.__len__())
    for li in fields_lis:
        print(li.text)

    field_name = 'Med Rec No'
    operator_value = 'Like'
    field_name_value = 'jon'

    field_xpath = '//div[@class="k-animation-container"][last()]/div/ul/li[text()="' + field_name + '"]'
    print(field_xpath)
    Smart_actions.wait_element_clickable(driver, By.XPATH, field_xpath)
    field_name_select = driver.find_element_by_xpath(field_xpath)
    field_name_select.click()

    # operator
    operator = driver.find_element_by_xpath(
        '//div[@id="customsearch-grid-div"]/div/div[@class="k-grid-content"]/table/tbody/tr[last()]/td[4]/span')
    operator.click()

    Smart_actions.wait_presence_element_xpath(driver,
                                              '//div[@class="k-animation-container"][last()]/div/ul/li[text()="' + operator_value + '"]')
    operator_selected = driver.find_element_by_xpath(
        '//div[@class="k-animation-container"][last()]/div/ul/li[text()="' + operator_value + '"]')
    operator_selected.click()

    # value
    value_input_xpath = './/div[@id="customsearch-grid"]/div[@class="k-grid-content"]/table/tbody/tr[last()]/td[5]/div/div[6]/input'
    value_input = driver.find_element_by_xpath(value_input_xpath)
    value_input.click()
    value_input.send_keys(field_name_value)

    reports.reports.view_report(driver)
    reports.reports.switch_to_report_view_page(driver, report_name)

def no_default_filter_handle(driver, report_name):
    # click "+" icon
    Smart_actions.wait_element_clickable(driver, By.ID, 'btnInsertClause')
    add_icon = driver.find_element_by_id('btnInsertClause')
    add_icon.click()

    # add filters besides the default ones
    driver.find_element_by_id('customsearch-grid-div')
    fields = driver.find_elements_by_xpath(
        './/div[@id="customsearch-grid-div"]/div/div[@class="k-grid-content"]/table/tbody/tr')

    # click the last row of the CS
    field_last = driver.find_element_by_xpath(
        './/div[@id="customsearch-grid-div"]/div/div[@class="k-grid-content"]/table/tbody/tr[last()]/td[3]/span')
    field_last.click()

    fields_lis = driver.find_elements_by_xpath('//div[@class="k-animation-container"][last()]/div/ul/li')
    print(fields_lis.__len__())
    for li in fields_lis:
        print(li.text)

    if report_name in ['Outpatient Flag Information by Facility','Outpatient Flag Information','APC Reimbursement Information','Coder Information','Inpatient Flag Information','Inpatient Flag Information by Facility']:
        field_name = fields_lis[0].text
        operator_value = '= Equal'
        field_name_value = '1/2/2019'
        value_input_xpath_ = './/div[@id="customsearch-grid"]/div[@class="k-grid-content"]/table/tbody/tr[last()]/td[5]/div/div[2]/span/span/input'
    else:
        field_name = 'Med Rec No'
        operator_value = 'Like'
        field_name_value = 'jon'
        value_input_xpath_ = './/div[@id="customsearch-grid"]/div[@class="k-grid-content"]/table/tbody/tr[last()]/td[5]/div/div[6]/input'

    field_xpath = '//div[@class="k-animation-container"][last()]/div/ul/li[text()="' + field_name + '"]'
    print(field_xpath)
    Smart_actions.wait_element_clickable(driver, By.XPATH, field_xpath)
    field_name_select = driver.find_element_by_xpath(field_xpath)
    field_name_select.click()

    # operator
    operator = driver.find_element_by_xpath(
        '//div[@id="customsearch-grid-div"]/div/div[@class="k-grid-content"]/table/tbody/tr[last()]/td[4]/span')
    operator.click()

    Smart_actions.wait_presence_element_xpath(driver,
                                              '//div[@class="k-animation-container"][last()]/div/ul/li[text()="' + operator_value + '"]')
    operator_selected = driver.find_element_by_xpath(
        '//div[@class="k-animation-container"][last()]/div/ul/li[text()="' + operator_value + '"]')
    operator_selected.click()

    # value
    if field_name in ['Med Rec No']:
        value_input_xpath = value_input_xpath_
        value_input = driver.find_element_by_xpath(value_input_xpath)
        value_input.click()
        value_input.send_keys(field_name_value)
    elif field_name in ['APC Code']:
        value_input_xpath = value_input_xpath_
        value_input = driver.find_element_by_xpath(value_input_xpath)
        value_input.click()
        value_input.send_keys('1')

    elif field_name in ['Flag Group - Primary','Flag Group']:
        # click search icon
        value_a_xpath = './/div[@id="customsearch-grid"]/div[@class="k-grid-content"]/table/tbody/tr[last()]/td[5]/div/div[4]/span/a'
        ACT.wait_element_clickable(driver, By.XPATH, value_a_xpath)
        ACT.wait_presence_element_xpath(driver, value_a_xpath)
        value_a = driver.find_element_by_xpath(value_a_xpath)
        value_a.click()

        ACT.wait_invisibility_of_element_located(driver)
        ACT.wait_presence_element_xpath(driver,
                                        '//div[@id="dvLookupuGrid"]/div[@class="k-grid-content"]/table/tbody/tr[1]/td[2]')
        value_select = driver.find_element_by_xpath(
            '//div[@id="dvLookupuGrid"]/div[@class="k-grid-content"]/table/tbody/tr[1]/td[2]')

        value_select.click()

def error_occurred_handle(driver, report_name):
    Smart_pass = True
    if driver.page_source.__contains__('Error Occurred'):
        report_tools.take_screenshot(driver, report_name + 'out of memory')
        out_of_memory = 'error_occurred_handle'
        print(out_of_memory)
        reports.reports.close_report_view_window(driver)
        reports.reports.close_report_cs_page(driver)
        Smart_pass = False
    else:
        pass
    return Smart_pass

def out_of_memory_handle(driver, report_name):
    if driver.page_source.__contains__('System.OutOfMemoryException'):
        report_tools.take_screenshot(driver, report_name + 'out of memory')
        out_of_memory = 'System.OutOfMemoryException'
        print(out_of_memory)
        reports.reports.close_report_view_window(driver)
        # 加一个filter
        # set up filters
        # click "+" icon
        Smart_actions.wait_element_clickable(driver, By.ID, 'btnInsertClause')
        add_icon = driver.find_element_by_id('btnInsertClause')
        add_icon.click()

        # add filters besides the default ones
        driver.find_element_by_id('customsearch-grid-div')
        fields = driver.find_elements_by_xpath(
            './/div[@id="customsearch-grid-div"]/div/div[@class="k-grid-content"]/table/tbody/tr')

        # click the last row of the CS
        field_last = driver.find_element_by_xpath(
            './/div[@id="customsearch-grid-div"]/div/div[@class="k-grid-content"]/table/tbody/tr[last()]/td[3]/span')
        field_last.click()

        fields_lis = driver.find_elements_by_xpath('//div[@class="k-animation-container"][last()]/div/ul/li')
        print(fields_lis.__len__())
        for li in fields_lis:
            print(li.text)

        if 'DRG Listing by Payer' in report_name:
            field_name = fields_lis[0].text
            operator_value = '> Greater Than'
            field_name_value = '1/2/2019'
            value_input_xpath_ = './/div[@id="customsearch-grid"]/div[@class="k-grid-content"]/table/tbody/tr[last()]/td[5]/div/div[2]/span/span/input'
        else:
            field_name = 'Med Rec No'
            operator_value = 'Like'
            field_name_value = 'jon'
            value_input_xpath_ = './/div[@id="customsearch-grid"]/div[@class="k-grid-content"]/table/tbody/tr[last()]/td[5]/div/div[6]/input'

        field_xpath = '//div[@class="k-animation-container"][last()]/div/ul/li[text()="' + field_name + '"]'
        print(field_xpath)
        Smart_actions.wait_element_clickable(driver, By.XPATH, field_xpath)
        field_name_select = driver.find_element_by_xpath(field_xpath)
        field_name_select.click()

        # operator
        operator = driver.find_element_by_xpath(
            '//div[@id="customsearch-grid-div"]/div/div[@class="k-grid-content"]/table/tbody/tr[last()]/td[4]/span')
        operator.click()

        Smart_actions.wait_presence_element_xpath(driver,
                                                  '//div[@class="k-animation-container"][last()]/div/ul/li[text()="' + operator_value + '"]')
        operator_selected = driver.find_element_by_xpath(
            '//div[@class="k-animation-container"][last()]/div/ul/li[text()="' + operator_value + '"]')
        operator_selected.click()

        # value
        value_input_xpath = value_input_xpath_
        value_input = driver.find_element_by_xpath(value_input_xpath)
        value_input.click()
        value_input.send_keys(field_name_value)

        reports.reports.view_report(driver, report_name)
        reports.reports.switch_to_report_view_page(driver, report_name)


def read_report_names(enterprise_or_standard):
    file_name = r'C:\Users\yyang212\PycharmProjects\PythonStudy\SMART\Smart_regression\resources\IPReport_names.txt'
    print('read_report_names:' + file_name)
    report_file = open(file_name, 'r')
    lines = report_file.readlines()

    names_enterprise = []
    names_standard = []

    print('strat to read file:' + file_name)
    for line in lines:
        print(line.strip())
        if not (line.strip().startswith('-') or line.strip().startswith('1')):
            if '-standard' in line:
                print('-----------Find IP standard report-------------')
                names_standard.append((line.strip()[0:-9]).strip())

            elif '-enterprise' in line:
                print('-----------FInd IP enterprise report-------------')
                names_enterprise.append((line.strip()[0:-11]).strip())
            else:
                print('nor standard or enterprise')
                pass

    report_file.close()

    if enterprise_or_standard == 'standard':
        return names_standard
    elif enterprise_or_standard == 'enterprise':
        return names_enterprise


def take_screenshot(driver, report_name):
    if report_name == 'Top 50 CC/MCC Diagnoses':
        picture_path = '../screenshots/Top 50 CC_MCC Diagnoses.png'
    else:
        picture_path = '../screenshots/' + report_name + '.png'

    # if pic exists, delete it
    tools.file_exist_delete(picture_path)
    driver.save_screenshot(picture_path)

def write_test_result_as_docx(report_name, test_case):

    if os.path.exists('../results/'+settings.test_result_file_name):
        document = Document('../results/'+settings.test_result_file_name)
    else:
        document = Document()
    document.add_heading(test_case, level=1)
    document.add_picture(report_name, width=Inches(6.43))
    document.save('../results/'+settings.test_result_file_name)

def write_test_result_as_docx_com(screenshot_path, heading_value):

    if os.path.exists('../results/'+settings.test_result_file_name):
        document = Document('../results/'+settings.test_result_file_name)
    else:
        document = Document()
    document.add_heading(heading_value, level=1)
    document.add_picture(screenshot_path, width=Inches(6.43))
    document.save('../results/'+settings.test_result_file_name)